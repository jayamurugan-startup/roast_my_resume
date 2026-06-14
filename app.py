"""
RoastMyResume — single-file Flask app (IMPROVED EDITION).

- Paste resume text OR upload a .docx / .txt file.
- Sends it to Groq (Llama 4 Scout, falling back to Llama 3.3) via plain `requests`.
- AI roasts the resume with hype-man / battle-rap swagger — funny and savage
  about the WRITING, never about the person.
- Frontend renders a "graded exam" style result + a square share-card that
  can be downloaded as a PNG (via html2canvas) for sharing.

Run locally:
    export GROQ_API_KEY=your_key_here
    pip install -r requirements.txt
    python app.py

Deploy on Render:
    Build command: pip install -r requirements.txt
    Start command: gunicorn app:app
    Env var:        GROQ_API_KEY=your_key_here
"""

import io
import requests
from flask import Flask, render_template_string, request, jsonify
from docx import Document

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 2 * 1024 * 1024  # 2 MB upload limit

# ---------------------------------------------------------------------------
# Groq config
# ---------------------------------------------------------------------------

# ⚠️ Hardcoded for quick deploy. If this repo is public, anyone can read this
# key and use your Groq quota. Rotate it at console.groq.com if that happens.
GROQ_API_KEY = "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"

GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
PRIMARY_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
FALLBACK_MODEL = "llama-3.3-70b-versatile"
MAX_CHARS = 8000

SYSTEM_PROMPT = """You are RoastMyResume — a savage battle-rap hype-man with the swagger of a WWE announcer and the 
honesty of your realest friend who has ZERO filter. You roast resumes for laughs, the kind of roast people screenshot 
and send to the group chat with "NO HE DIDN'T 💀" and "I'M SCREAMING 😂"

Your roasting style:
- Be BRUTALLY SPECIFIC — quote or paraphrase EXACT lines/words from the resume and translate them into what they ACTUALLY mean. 
  For example: "7 years of experience designing, building, and deploying robust full-stack applications" becomes 
  "7 years of experience making stuff up as I go along and hoping nobody notices 🎭"
- Use the "What They Wrote vs What It Actually Means" format for bullet points — this is your SECRET SAUCE.
- Big confident energy and playful trash talk ("we got a heavyweight buzzword champion in the building 🏆"), 
  but NEVER cruel. Roast the WRITING, the buzzwords, the formatting, the vague bullet points — never the person's 
  worth, identity, age, employment gaps, or anything they can't control.
- Use emojis liberally and on-theme (🔥💯🕶️🐐👑💀🏆🎤🎭📉🚩).
- Close on a hype note — like you're sending them into the ring ready to win the job.

Respond in EXACTLY this format, with these exact emoji headers, and nothing before or after:

🔥 ROAST SCORE: X/10

💀 THE VERDICT
(2-3 punchy sentences, overall savage-but-funny take. Make it quotable. Make it sting.)

🎯 THE RECEIPTS (What They Wrote vs What It Actually Means)
(4-6 bullet points. Each bullet must follow this EXACT format:
- "[Exact or paraphrased quote from resume]" → "[What it actually means / savage translation]" 
Example: - "Led cross-functional teams to deliver scalable solutions" → "Sent some Slack messages and prayed the interns figured it out" 💀
)

😬 BUZZWORD ALERT
(2-3 sentences calling out the most overused corporate buzzwords found, with mock translations of what they actually mean. 
Be specific about WHICH buzzwords you found.)

✨ RESPECT (BEGRUDGING)
(1 sentence, genuinely find something that's actually good, said reluctantly and with maximum sass)

🐦 QUOTE OF THE DAY
(ONE single savage-but-funny one-liner, under 200 characters, that sums up this resume — 
written so it could be screenshotted and shared on its own. Make it HURT but make it FUNNY.)
"""


def call_groq(model: str, resume_text: str) -> str:
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"Roast this resume with MAXIMUM savagery but keep it about the writing:\n\n{resume_text}"},
        ],
        "temperature": 0.95,
        "max_tokens": 900,
    }
    resp = requests.post(GROQ_API_URL, headers=headers, json=payload, timeout=30)
    resp.raise_for_status()
    data = resp.json()
    return data["choices"][0]["message"]["content"]


def extract_text_from_docx(file_stream) -> str:
    document = Document(file_stream)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # also grab text inside tables (lots of resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template_string(INDEX_HTML)


@app.route("/roast", methods=["POST"])
def roast():
    resume_text = ""

    uploaded = request.files.get("file")
    if uploaded and uploaded.filename:
        filename = uploaded.filename.lower()
        try:
            if filename.endswith(".docx"):
                resume_text = extract_text_from_docx(io.BytesIO(uploaded.read()))
            elif filename.endswith(".txt"):
                resume_text = uploaded.read().decode("utf-8", errors="ignore")
            else:
                return jsonify(
                    {"error": "Only .docx or .txt, champ. Or just paste the text 📋"}
                ), 400
        except Exception:
            return jsonify(
                {"error": "Couldn't read that file. Re-save it and try again 📄"}
            ), 400
    else:
        resume_text = (request.form.get("resume") or "").strip()

    resume_text = resume_text.strip()

    if not resume_text:
        return jsonify({"error": "Paste some resume text or upload a file first 📄"}), 400

    if len(resume_text) < 30:
        return jsonify({"error": "That's not a resume, that's a haiku. Give me more 📝"}), 400

    resume_text = resume_text[:MAX_CHARS]

    try:
        roast_text = call_groq(PRIMARY_MODEL, resume_text)
    except Exception:
        try:
            roast_text = call_groq(FALLBACK_MODEL, resume_text)
        except Exception:
            return jsonify(
                {"error": "Even the roast bot rage-quit reading this. Try again in a sec 💀"}
            ), 500

    return jsonify({"roast": roast_text})


# ---------------------------------------------------------------------------
# Frontend (single template, inline CSS + JS)
# ---------------------------------------------------------------------------

INDEX_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>RoastMyResume 🕶️🔥 — AI roasts your resume (with love)</title>
<meta name="description" content="Paste or upload your resume. Get roasted with swagger. Download the card. Share the burn.">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Caveat:wght@600;700&family=Courier+Prime:ital,wght@0,400;0,700;1,400&family=Inter:wght@500;600;700;800;900&display=swap" rel="stylesheet">
<script src="https://cdnjs.cloudflare.com/ajax/libs/html2canvas/1.4.1/html2canvas.min.js"></script>
<style>
  :root {
    --desk: #1a1a2e;
    --desk-accent: #16213e;
    --paper: #FBF6EC;
    --paper-edge: #EFE6D2;
    --paper-line: #E3D9C4;
    --ink: #2A2622;
    --ink-soft: #5b554c;
    --red: #E63946;
    --red-dark: #c1121f;
    --red-glow: rgba(230, 57, 70, 0.3);
    --gold: #F2B705;
    --gold-light: #FFD700;
    --green-pen: #2F7D4F;
    --dark: #0f0f1a;
    --card-bg: #16161A;
    --accent: #E63946;
  }

  * { box-sizing: border-box; margin: 0; padding: 0; }

  html, body { min-height: 100%; }

  body {
    font-family: 'Inter', sans-serif;
    background: var(--desk);
    background-image:
      radial-gradient(circle at 15% 15%, rgba(230, 57, 70, 0.08), transparent 50%),
      radial-gradient(circle at 85% 85%, rgba(242, 183, 5, 0.06), transparent 50%),
      radial-gradient(circle at 50% 50%, rgba(22, 33, 62, 0.5), transparent 70%);
    color: var(--ink);
    min-height: 100vh;
    display: flex;
    justify-content: center;
    padding: clamp(20px, 5vw, 56px) 16px 80px;
  }

  .wrap { width: 100%; max-width: 800px; }

  .site-header { text-align: center; color: var(--paper); margin-bottom: 32px; }

  .site-header .brand {
    font-family: 'Caveat', cursive;
    font-weight: 700;
    font-size: clamp(2.8rem, 10vw, 4.8rem);
    line-height: 1;
    color: var(--gold);
    transform: rotate(-2deg);
    display: inline-block;
    text-shadow: 3px 3px 0 rgba(0,0,0,0.4), 0 0 40px rgba(242, 183, 5, 0.2);
    animation: float 3s ease-in-out infinite;
  }

  @keyframes float {
    0%, 100% { transform: rotate(-2deg) translateY(0); }
    50% { transform: rotate(-2deg) translateY(-6px); }
  }

  .site-header .tagline {
    margin: 14px 0 0;
    font-size: 1rem;
    color: rgba(251,246,236,0.7);
    font-weight: 500;
  }

  .site-header .sub-tagline {
    margin: 6px 0 0;
    font-size: 0.85rem;
    color: rgba(251,246,236,0.45);
    font-family: 'Courier Prime', monospace;
  }

  /* ---------- Paper card ---------- */
  .paper {
    background: var(--paper);
    border-radius: 8px;
    padding: clamp(24px, 5vw, 48px);
    box-shadow: 
      0 1px 0 var(--paper-edge), 
      0 20px 60px -15px rgba(0,0,0,0.5), 
      0 0 0 1px rgba(0,0,0,0.04),
      0 0 80px rgba(230, 57, 70, 0.08);
    position: relative;
    transform: rotate(-0.3deg);
    transition: transform 0.3s ease;
  }

  .paper:hover { transform: rotate(0deg) scale(1.002); }

  .paper::before {
    content: "";
    position: absolute;
    inset: 0;
    border-radius: 8px;
    background-image: repeating-linear-gradient(to bottom, transparent, transparent 37px, var(--paper-line) 38px);
    opacity: 0.5;
    pointer-events: none;
  }

  .paper-content { position: relative; z-index: 1; }

  .paper-head {
    font-family: 'Courier Prime', monospace;
    font-weight: 700;
    font-size: 0.78rem;
    letter-spacing: 0.2em;
    text-transform: uppercase;
    color: var(--ink-soft);
    margin: 0 0 18px;
    padding-bottom: 14px;
    border-bottom: 2px solid var(--ink);
    display: flex;
    justify-content: space-between;
    align-items: baseline;
    gap: 12px;
  }

  .paper-head span:last-child { font-weight: 400; letter-spacing: 0.06em; color: #b9b1a2; }

  /* Tabs */
  .tabs { display: flex; gap: 10px; margin-bottom: 16px; }

  .tab-btn {
    font-family: 'Inter', sans-serif;
    font-weight: 700;
    font-size: 0.88rem;
    padding: 8px 20px;
    border-radius: 999px;
    border: 2.5px solid #000000;
    background: transparent;
    color: #000000;
    cursor: pointer;
    transition: all 0.2s ease;
    position: relative;
    overflow: hidden;
  }

  .tab-btn::before {
    content: '';
    position: absolute;
    inset: 0;
    background: var(--ink);
    transform: scaleX(0);
    transform-origin: left;
    transition: transform 0.2s ease;
    z-index: -1;
  }

  .tab-btn:hover::before { transform: scaleX(1); }
  .tab-btn:hover { color: var(--paper); }

  .tab-btn.active { background: var(--ink); color: var(--paper); }
  .tab-btn.active::before { transform: scaleX(1); }

  textarea#resume {
    width: 100%;
    min-height: 260px;
    resize: vertical;
    border: none;
    background: transparent;
    font-family: 'Courier Prime', monospace;
    font-size: 1rem;
    line-height: 38px;
    color: var(--ink);
    padding: 0;
    outline: none;
  }

  textarea#resume::placeholder { color: #b6ad9c; font-style: italic; }

  /* Upload zone */
  .upload-zone {
    display: none;
    border: 2.5px dashed var(--ink-soft);
    border-radius: 12px;
    padding: 50px 24px;
    text-align: center;
    font-family: 'Courier Prime', monospace;
    color: var(--ink-soft);
    cursor: pointer;
    transition: all 0.2s ease;
    background: rgba(0,0,0,0.02);
  }

  .upload-zone.active { display: block; }
  .upload-zone:hover, .upload-zone.dragover { 
    border-color: var(--red); 
    color: var(--red-dark); 
    background: rgba(230, 57, 70, 0.04);
    transform: scale(1.01);
  }
  .upload-zone input { display: none; }
  .upload-zone .filename { margin-top: 12px; font-weight: 700; color: var(--ink); font-size: 1.1rem; }
  .upload-zone .file-icon { font-size: 2.5rem; margin-bottom: 8px; display: block; }

  /* Action button */
  .action-row { display: flex; justify-content: center; margin: 32px 0 8px; }

  .stamp-btn {
    font-family: 'Caveat', cursive;
    font-weight: 700;
    font-size: 1.9rem;
    color: var(--red);
    background: transparent;
    border: 4px solid var(--red);
    border-radius: 12px;
    padding: 12px 48px;
    cursor: pointer;
    transform: rotate(-3deg);
    transition: all 0.2s ease;
    position: relative;
    overflow: hidden;
    box-shadow: 0 4px 20px rgba(230, 57, 70, 0.2);
  }

  .stamp-btn::before {
    content: '';
    position: absolute;
    inset: 0;
    background: var(--red);
    transform: scaleX(0);
    transform-origin: left;
    transition: transform 0.25s ease;
    z-index: -1;
  }

  .stamp-btn:hover { 
    transform: rotate(-3deg) scale(1.06); 
    color: var(--paper);
    box-shadow: 0 8px 30px rgba(230, 57, 70, 0.4);
  }
  .stamp-btn:hover::before { transform: scaleX(1); }

  .stamp-btn:active { transform: rotate(-1deg) scale(0.96); }
  .stamp-btn:disabled { opacity: 0.5; cursor: wait; }
  .stamp-btn.loading { background: var(--red); color: var(--paper); }
  .stamp-btn.loading::before { transform: scaleX(1); }

  .hint { text-align: center; font-size: 0.82rem; color: var(--ink-soft); margin-top: 8px; font-family: 'Courier Prime', monospace; }

  .error-msg {
    text-align: center;
    color: var(--red-dark);
    font-family: 'Caveat', cursive;
    font-weight: 700;
    font-size: 1.4rem;
    margin-top: 16px;
    min-height: 1.8rem;
    animation: shake 0.5s ease;
  }

  @keyframes shake {
    0%, 100% { transform: translateX(0); }
    25% { transform: translateX(-8px); }
    75% { transform: translateX(8px); }
  }

  /* ---------- Results ---------- */
  #results { margin-top: 40px; display: none; }
  #results.visible { display: block; animation: slideUp 0.6s ease; }

  @keyframes slideUp {
    from { opacity: 0; transform: translateY(30px); }
    to { opacity: 1; transform: translateY(0); }
  }

  .score-wrap { display: flex; justify-content: center; margin-bottom: 24px; }

  .score-circle { 
    position: relative; 
    width: 170px; 
    height: 170px; 
    filter: drop-shadow(0 8px 20px rgba(230, 57, 70, 0.3));
  }
  .score-circle svg { width: 100%; height: 100%; transform: rotate(-7deg); }

  .score-circle path {
    fill: none;
    stroke: var(--red);
    stroke-width: 7;
    stroke-linecap: round;
    stroke-dasharray: 600;
    stroke-dashoffset: 600;
    transition: stroke-dashoffset 1.2s cubic-bezier(0.34, 1.56, 0.64, 1);
  }

  .score-circle .score-text {
    position: absolute; inset: 0;
    display: flex; align-items: center; justify-content: center; flex-direction: column;
    font-family: 'Caveat', cursive; font-weight: 700; color: var(--red);
  }

  .score-circle .score-num { font-size: 3.6rem; line-height: 1; text-shadow: 2px 2px 0 rgba(0,0,0,0.1); }
  .score-circle .score-out { font-size: 1.1rem; color: var(--ink-soft); margin-top: -2px; }

  .note {
    background: linear-gradient(135deg, var(--gold) 0%, #f0c040 100%);
    border-radius: 4px;
    padding: 18px 22px;
    margin: 0 0 26px;
    box-shadow: 0 8px 24px -8px rgba(0,0,0,0.35);
    transform: rotate(-1deg);
    position: relative;
    overflow: hidden;
  }

  .note::before {
    content: '';
    position: absolute;
    top: -20px; right: -20px;
    width: 80px; height: 80px;
    background: rgba(255,255,255,0.15);
    border-radius: 50%;
  }

  .note h3 { margin: 0 0 10px; font-family: 'Caveat', cursive; font-weight: 700; font-size: 1.7rem; color: var(--ink); }
  .note ul { margin: 0; padding-left: 1.2em; }
  .note li { font-family: 'Courier Prime', monospace; font-size: 0.95rem; line-height: 1.6; color: var(--ink); font-weight: 600; }
  .note ul li + li { margin-top: 10px; }

  .marker-card { padding: 10px 0 22px; }

  .marker-card h3 {
    font-family: 'Caveat', cursive; font-weight: 700; font-size: 1.8rem;
    color: var(--red); margin: 0 0 8px;
  }

  .marker-card.green h3 { color: var(--green-pen); }

  .marker-card .underline { display: block; width: 100%; max-width: 280px; height: 10px; margin: -6px 0 12px; }
  .marker-card .underline path { stroke: var(--red); stroke-width: 6; fill: none; stroke-linecap: round; }
  .marker-card.green .underline path { stroke: var(--green-pen); }

  .marker-card p { font-family: 'Courier Prime', monospace; font-size: 1rem; line-height: 1.65; color: #ffffff; margin: 0; }

  /* Receipts special styling */
  .receipt-item {
    background: rgba(0,0,0,0.03);
    border-radius: 8px;
    padding: 14px 16px;
    margin-bottom: 12px;
    border-left: 4px solid var(--red);
  }

  .receipt-item .wrote {
    font-family: 'Courier Prime', monospace;
    font-size: 0.9rem;
    color: var(--ink-soft);
    font-style: italic;
    margin-bottom: 6px;
  }

  .receipt-item .means {
    font-family: 'Inter', sans-serif;
    font-weight: 700;
    font-size: 1rem;
    color: var(--red-dark);
  }

  .receipt-item .arrow {
    color: var(--gold);
    font-weight: 900;
    margin: 0 6px;
  }

  /* Tweetable burn */
  .burn-box {
    background: linear-gradient(135deg, var(--dark) 0%, #1a1a2e 100%);
    color: var(--paper);
    border-radius: 12px;
    padding: 24px 26px;
    margin-top: 12px;
    transform: rotate(0.5deg);
    position: relative;
    overflow: hidden;
    border: 1px solid rgba(230, 57, 70, 0.2);
  }

  .burn-box::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 4px;
    background: linear-gradient(90deg, var(--red), var(--gold), var(--red));
  }

  .burn-box .label { font-family: 'Caveat', cursive; font-size: 1.6rem; color: var(--gold); font-weight: 700; margin-bottom: 10px; }
  .burn-box .text { 
    font-family: 'Courier Prime', monospace; 
    font-size: 1.1rem; 
    line-height: 1.6; 
    font-weight: 700;
    color: var(--paper);
  }

  .burn-actions { display: flex; gap: 12px; margin-top: 20px; flex-wrap: wrap; }

  .burn-actions button, .burn-actions a {
    font-family: 'Inter', sans-serif; font-weight: 700; font-size: 0.88rem;
    border-radius: 8px; padding: 10px 18px; border: 1.5px solid rgba(251,246,236,0.25);
    background: transparent; color: var(--paper); cursor: pointer; text-decoration: none;
    transition: all 0.2s ease;
    display: inline-flex;
    align-items: center;
    gap: 6px;
  }

  .burn-actions button:hover, .burn-actions a:hover { 
    background: rgba(251,246,236,0.12); 
    border-color: var(--gold);
    transform: translateY(-2px);
  }

  .again-row { text-align: center; margin-top: 32px; }

  .again-row button {
    font-family: 'Caveat', cursive; font-weight: 700; font-size: 1.3rem;
    color: var(--ink-soft); background: transparent; border: none; cursor: pointer; text-decoration: underline;
    transition: color 0.2s ease;
  }
  .again-row button:hover { color: var(--red); }

  /* ---------- Share card (captured by html2canvas) ---------- */
  .share-card-outer { margin-top: 36px; display: flex; justify-content: center; }

  .share-card {
    width: 360px;
    aspect-ratio: 1 / 1;
    background: linear-gradient(145deg, #1e1e3a 0%, #2a2a4a 50%, #1e1e3a 100%);
    border-radius: 24px;
    padding: 2px;
    display: flex;
    flex-direction: column;
    justify-content: space-between;
    color: #FBF6EC;
    font-family: 'Inter', sans-serif;
    position: relative;
    overflow: hidden;
    border: 2px solid rgba(242, 183, 5, 0.3);
    box-shadow: 
      0 20px 60px rgba(0,0,0,0.5),
      0 0 40px rgba(230, 57, 70, 0.15),
      inset 0 1px 0 rgba(255,255,255,0.08);
  }

  .share-card::before {
    content: "";
    position: absolute;
    top: -80px; right: -80px;
    width: 250px; height: 250px;
    background: radial-gradient(circle, rgba(242,183,5,0.25), transparent 70%);
    pointer-events: none;
  }

  .share-card::after {
    content: "";
    position: absolute;
    bottom: -60px; left: -60px;
    width: 180px; height: 180px;
    background: radial-gradient(circle, rgba(230,57,70,0.2), transparent 70%);
    pointer-events: none;
  }

  .share-card .card-glow {
    position: absolute;
    top: 50%; left: 50%;
    transform: translate(-50%, -50%);
    width: 200px; height: 200px;
    background: radial-gradient(circle, rgba(242,183,5,0.08), transparent 70%);
    pointer-events: none;
    z-index: 0;
  }

  .share-card-header {
    display: flex; align-items: center; gap: 10px;
    font-family: 'Caveat', cursive; font-weight: 700; font-size: 1.6rem; color: #F2B705;
    position: relative; z-index: 1;
    text-shadow: 0 2px 4px rgba(0,0,0,0.3);
  }

  .share-card-score {
    font-family: 'Caveat', cursive; font-weight: 700; font-size: 6rem; color: #F2B705;
    line-height: 1; text-align: center;
    position: relative; z-index: 1;
    text-shadow: 0 0 40px rgba(242, 183, 5, 0.4), 0 2px 4px rgba(0,0,0,0.3);
  }

  .share-card-score span.out { font-size: 2rem; color: rgba(251,246,236,0.5); }

  .share-card-quote {
    font-family: 'Courier Prime', monospace;
    font-size: 1.05rem;
    line-height: 1.55;
    text-align: center;
    position: relative; z-index: 1;
    font-weight: 700;
    color: #FBF6EC;
    text-shadow: 0 2px 8px rgba(0,0,0,0.4);
    padding: 0 8px;
  }

  .share-card-footer {
    font-family: 'Inter', sans-serif; font-weight: 700; font-size: 0.75rem;
    letter-spacing: 0.18em; text-transform: uppercase; text-align: center;
    color: rgba(251,246,236,0.55);
    position: relative; z-index: 1;
  }

  .share-card .flame {
    position: absolute;
    font-size: 3rem;
    opacity: 0.15;
    z-index: 0;
  }
  .share-card .flame-1 { top: 20px; right: 30px; }
  .share-card .flame-2 { bottom: 80px; left: 20px; transform: rotate(-20deg); }

  .download-row { text-align: center; margin-top: 18px; }

  .download-row button {
    font-family: 'Inter', sans-serif; font-weight: 700; font-size: 0.9rem;
    background: linear-gradient(135deg, var(--gold) 0%, #f0c040 100%); 
    color: var(--ink); border: none; border-radius: 8px;
    padding: 12px 28px; cursor: pointer;
    transition: all 0.2s ease;
    box-shadow: 0 4px 15px rgba(242, 183, 5, 0.3);
  }
  .download-row button:hover { 
    transform: translateY(-3px); 
    box-shadow: 0 8px 25px rgba(242, 183, 5, 0.4);
  }

  /* Loading animation */
  .loading-dots {
    display: inline-flex;
    gap: 4px;
  }
  .loading-dots span {
    width: 8px; height: 8px;
    background: var(--paper);
    border-radius: 50%;
    animation: bounce 0.6s ease infinite;
  }
  .loading-dots span:nth-child(2) { animation-delay: 0.1s; }
  .loading-dots span:nth-child(3) { animation-delay: 0.2s; }

  @keyframes bounce {
    0%, 100% { transform: translateY(0); }
    50% { transform: translateY(-8px); }
  }

  footer { text-align: center; color: rgba(251,246,236,0.35); font-size: 0.8rem; margin-top: 40px; font-family: 'Courier Prime', monospace; }

  @media (max-width: 480px) {
    .paper { transform: none; }
    .paper:hover { transform: none; }
    textarea#resume { line-height: 32px; min-height: 200px; }
    .share-card { width: 100%; max-width: 340px; padding: 24px; }
    .share-card-score { font-size: 5rem; }
    .stamp-btn { font-size: 1.5rem; padding: 10px 32px; }
  }

  @media (prefers-reduced-motion: reduce) {
    .score-circle path { transition: none; }
    .site-header .brand { animation: none; }
    #results.visible { animation: none; }
  }
</style>
</head>
<body>
  <div class="wrap">
    <header class="site-header">
      <div class="brand">RoastMyResume 🕶️🔥</div>
      <p class="tagline">Paste it. Upload it. We clown it. You screenshot it.</p>
      <p class="sub-tagline">"7 years of experience designing, building, and deploying robust full-stack applications"<br>→ "7 years of making stuff up as I go along and hoping nobody notices"</p>
    </header>

    <div class="paper">
      <div class="paper-content">
        <div class="paper-head">
          <span>Step Into The Ring</span>
          <span id="charcount">0 / 8000</span>
        </div>

        <div class="tabs">
          <button class="tab-btn active" id="tabPaste">📝 Paste text</button>
          <button class="tab-btn" id="tabUpload">📎 Upload .docx / .txt</button>
        </div>

        <textarea id="resume" placeholder="Paste your full resume text here — work history, bullet points, skills, all of it. The more buzzwords, the spicier the roast.&#10;&#10;Example of what we do:&#10;&quot;Led cross-functional teams to deliver scalable solutions&quot; → &quot;Sent some Slack messages and prayed the interns figured it out&quot; 💀" spellcheck="false"></textarea>

        <div class="upload-zone" id="uploadZone">
          <input type="file" id="fileInput" accept=".docx,.txt">
          <span class="file-icon">📄</span>
          Click or drag your .docx / .txt resume here
          <div class="filename" id="filename"></div>
        </div>

        <div class="action-row">
          <button id="roastBtn" class="stamp-btn">ROAST ME 🔥</button>
        </div>
        <p class="hint">Free. Anonymous. Mildly traumatic. Nothing is stored.</p>
        <div class="error-msg" id="error"></div>
      </div>
    </div>

    <div id="results">
      <div class="score-wrap">
        <div class="score-circle">
          <svg viewBox="0 0 200 200">
            <path id="scorePath" d="M100,10 C 150,10 190,50 190,100 C 190,150 150,190 100,190 C 50,190 10,150 10,100 C 10,50 50,10 100,10 Z" />
          </svg>
          <div class="score-text">
            <div class="score-num" id="scoreNum">?</div>
            <div class="score-out">/ 10</div>
          </div>
        </div>
      </div>

      <div class="marker-card">
        <h3>💀 The Verdict</h3>
        <svg class="underline" viewBox="0 0 240 8" preserveAspectRatio="none"><path d="M2,5 C60,2 120,7 238,3" /></svg>
        <p id="verdict"></p>
      </div>

      <div class="note">
        <h3>🎯 The Receipts</h3>
        <div id="receipts"></div>
      </div>

      <div class="marker-card">
        <h3>😬 Buzzword Alert</h3>
        <svg class="underline" viewBox="0 0 240 8" preserveAspectRatio="none"><path d="M2,5 C60,2 120,7 238,3" /></svg>
        <p id="buzzword"></p>
      </div>

      <div class="marker-card green">
        <h3>✨ Respect (Begrudging)</h3>
        <svg class="underline" viewBox="0 0 240 8" preserveAspectRatio="none"><path d="M2,5 C60,2 120,7 238,3" /></svg>
        <p id="compliment"></p>
      </div>

      <div class="burn-box">
        <div class="label">🐦 Quote Of The Day</div>
        <div class="text" id="tweetable"></div>
        <div class="burn-actions">
          <button id="copyBtn">📋 Copy burn</button>
          <a id="shareBtn" href="#" target="_blank" rel="noopener">🐦 Share on X →</a>
        </div>
      </div>

      <div class="share-card-outer">
        <div class="share-card" id="shareCard">
          <div class="card-glow"></div>
          <span class="flame flame-1">🔥</span>
          <span class="flame flame-2">💀</span>
          <div class="share-card-header">🕶️🔥 RoastMyResume</div>
          <div class="share-card-score"><span id="shareScoreNum">?</span><span class="out">/10</span></div>
          <div class="share-card-quote" id="shareQuote"></div>
          <div class="share-card-footer">roastmyresume — built different</div>
        </div>
      </div>
      <div class="download-row">
        <button id="downloadBtn">📸 Download card as image</button>
      </div>

      <div class="again-row">
        <button id="againBtn">🔁 roast another resume</button>
      </div>
    </div>

    <footer>
      Build for Product Hunt
    </footer>
  </div>

<script>
  const resumeEl = document.getElementById('resume');
  const charcount = document.getElementById('charcount');
  const roastBtn = document.getElementById('roastBtn');
  const errorEl = document.getElementById('error');
  const results = document.getElementById('results');

  const tabPaste = document.getElementById('tabPaste');
  const tabUpload = document.getElementById('tabUpload');
  const uploadZone = document.getElementById('uploadZone');
  const fileInput = document.getElementById('fileInput');
  const filenameEl = document.getElementById('filename');

  const scoreNum = document.getElementById('scoreNum');
  const scorePath = document.getElementById('scorePath');
  const verdictEl = document.getElementById('verdict');
  const receiptsEl = document.getElementById('receipts');
  const buzzwordEl = document.getElementById('buzzword');
  const complimentEl = document.getElementById('compliment');
  const tweetableEl = document.getElementById('tweetable');
  const copyBtn = document.getElementById('copyBtn');
  const shareBtn = document.getElementById('shareBtn');
  const againBtn = document.getElementById('againBtn');
  const downloadBtn = document.getElementById('downloadBtn');
  const shareScoreNum = document.getElementById('shareScoreNum');
  const shareQuote = document.getElementById('shareQuote');
  const shareCard = document.getElementById('shareCard');

  const MAX_CHARS = 8000;
  let mode = 'paste';

  // --- Tabs ---
  tabPaste.addEventListener('click', () => {
    mode = 'paste';
    tabPaste.classList.add('active');
    tabUpload.classList.remove('active');
    resumeEl.style.display = 'block';
    uploadZone.classList.remove('active');
    errorEl.textContent = '';
  });

  tabUpload.addEventListener('click', () => {
    mode = 'upload';
    tabUpload.classList.add('active');
    tabPaste.classList.remove('active');
    resumeEl.style.display = 'none';
    uploadZone.classList.add('active');
    errorEl.textContent = '';
  });

  uploadZone.addEventListener('click', () => fileInput.click());

  uploadZone.addEventListener('dragover', (e) => { e.preventDefault(); uploadZone.classList.add('dragover'); });
  uploadZone.addEventListener('dragleave', () => uploadZone.classList.remove('dragover'));
  uploadZone.addEventListener('drop', (e) => {
    e.preventDefault();
    uploadZone.classList.remove('dragover');
    if (e.dataTransfer.files.length) {
      fileInput.files = e.dataTransfer.files;
      filenameEl.textContent = fileInput.files[0].name;
    }
  });

  fileInput.addEventListener('change', () => {
    if (fileInput.files.length) filenameEl.textContent = fileInput.files[0].name;
  });

  resumeEl.addEventListener('input', () => {
    const len = resumeEl.value.length;
    charcount.textContent = `${Math.min(len, MAX_CHARS)} / ${MAX_CHARS}`;
    if (len > MAX_CHARS) {
      charcount.style.color = 'var(--red)';
    } else {
      charcount.style.color = '';
    }
  });

  // --- Parse roast text ---
  function parseRoast(text) {
    const data = { score: '?', verdict: [], receipts: [], buzzword: [], compliment: [], tweet: [] };
    let current = null;

    for (let raw of text.split('\n')) {
      const line = raw.trim();
      if (!line) continue;

      if (line.startsWith('🔥')) {
        const m = line.match(/(\d+(\.\d+)?)\s*\/\s*10/);
        data.score = m ? m[1] : '?';
        current = null;
        continue;
      }
      if (line.startsWith('💀')) { current = 'verdict'; continue; }
      if (line.startsWith('🎯')) { current = 'receipts'; continue; }
      if (line.startsWith('😬')) { current = 'buzzword'; continue; }
      if (line.startsWith('✨')) { current = 'compliment'; continue; }
      if (line.startsWith('🐦')) { current = 'tweet'; continue; }

      if (current) data[current].push(line.replace(/^[-•]\s*/, ''));
    }
    return data;
  }

  function render(data) {
    const scoreVal = parseFloat(data.score);
    scoreNum.textContent = isNaN(scoreVal) ? '?' : data.score;
    shareScoreNum.textContent = isNaN(scoreVal) ? '?' : data.score;

    if (!isNaN(scoreVal)) {
      const pathLen = 600;
      const offset = pathLen - (pathLen * (scoreVal / 10));
      scorePath.style.strokeDashoffset = pathLen;
      requestAnimationFrame(() => { scorePath.style.strokeDashoffset = offset; });
    }

    verdictEl.textContent = data.verdict.join(' ');

    // Render receipts with special formatting for "What They Wrote vs What It Means"
    receiptsEl.innerHTML = '';
    if (data.receipts.length) {
      data.receipts.forEach(item => {
        const div = document.createElement('div');
        div.className = 'receipt-item';

        // Try to parse "wrote" → "means" format
        const arrowMatch = item.match(/^(.*?)(?:→|->|—>|–>)(.*)$/);
        if (arrowMatch) {
          const wrote = arrowMatch[1].trim().replace(/^["']|["']$/g, '');
          const means = arrowMatch[2].trim().replace(/^["']|["']$/g, '');
          div.innerHTML = `
            <div class="wrote">&ldquo;${wrote}&rdquo;</div>
            <div class="means"><span class="arrow">→</span> ${means}</div>
          `;
        } else {
          div.innerHTML = `<div class="means">${item}</div>`;
        }
        receiptsEl.appendChild(div);
      });
    } else {
      receiptsEl.innerHTML = '<div class="receipt-item"><div class="means">Honestly, it speaks for itself. 💀</div></div>';
    }

    buzzwordEl.textContent = data.buzzword.join(' ') || "Surprisingly buzzword-light. We're shook.";
    complimentEl.textContent = data.compliment.join(' ') || 'It exists. That counts for something.';

    const tweet = data.tweet.join(' ') || data.verdict.join(' ').slice(0, 200);
    tweetableEl.textContent = tweet;
    shareQuote.textContent = tweet;

    const shareText = encodeURIComponent(`${tweet} 🔥 — roasted by RoastMyResume`);
    shareBtn.href = `https://twitter.com/intent/tweet?text=${shareText}`;

    results.classList.add('visible');
    results.scrollIntoView({ behavior: 'smooth', block: 'start' });
  }

  // --- Submit ---
  async function roast() {
    errorEl.textContent = '';
    const formData = new FormData();

    if (mode === 'upload') {
      if (!fileInput.files.length) {
        errorEl.textContent = 'Drop a .docx or .txt file first 📄';
        return;
      }
      formData.append('file', fileInput.files[0]);
    } else {
      const resumeText = resumeEl.value.trim();
      if (resumeText.length < 30) {
        errorEl.textContent = "That's not a resume, that's a haiku. Give me more 📝";
        return;
      }
      formData.append('resume', resumeText);
    }

    roastBtn.disabled = true;
    roastBtn.classList.add('loading');
    roastBtn.innerHTML = 'STEPPING INTO THE RING<span class="loading-dots"><span></span><span></span><span></span></span>';

    try {
      const res = await fetch('/roast', { method: 'POST', body: formData });
      const payload = await res.json();

      if (!res.ok) {
        errorEl.textContent = payload.error || 'Something broke. Try again.';
        return;
      }

      render(parseRoast(payload.roast));
    } catch (err) {
      errorEl.textContent = 'Network gremlins. Try again in a sec.';
    } finally {
      roastBtn.disabled = false;
      roastBtn.classList.remove('loading');
      roastBtn.textContent = 'ROAST ME 🔥';
    }
  }

  roastBtn.addEventListener('click', roast);

  copyBtn.addEventListener('click', () => {
    navigator.clipboard.writeText(tweetableEl.textContent).then(() => {
      copyBtn.textContent = '✅ Copied!';
      setTimeout(() => copyBtn.textContent = '📋 Copy burn', 1500);
    });
  });

  downloadBtn.addEventListener('click', () => {
    downloadBtn.textContent = '⏳ Generating...';
    html2canvas(shareCard, { scale: 3, backgroundColor: null }).then(canvas => {
      const link = document.createElement('a');
      link.download = 'roastmyresume.png';
      link.href = canvas.toDataURL('image/png');
      link.click();
      downloadBtn.textContent = '📸 Download card as image';
    }).catch(() => {
      downloadBtn.textContent = '❌ Failed. Try again.';
      setTimeout(() => downloadBtn.textContent = '📸 Download card as image', 2000);
    });
  });

  againBtn.addEventListener('click', () => {
    results.classList.remove('visible');
    resumeEl.value = '';
    fileInput.value = '';
    filenameEl.textContent = '';
    charcount.textContent = `0 / ${MAX_CHARS}`;
    window.scrollTo({ top: 0, behavior: 'smooth' });
  });
</script>
</body>
</html>
"""


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)